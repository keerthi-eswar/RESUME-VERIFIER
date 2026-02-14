"""
Document Processor Module
Handles PDF and DOCX text extraction
"""

import io
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
import re
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles text extraction from PDF and DOCX files"""

    def __init__(self):
        """Initialize document processor"""
        self.stopwords = self._load_stopwords()

    def _load_stopwords(self) -> set:
        """Load common English stopwords"""
        common_stopwords = {
            'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from',
            'has', 'he', 'in', 'is', 'it', 'its', 'of', 'on', 'or', 'that',
            'the', 'to', 'was', 'will', 'with', 'i', 'me', 'my', 'we', 'you',
            'your', 'this', 'but', 'not', 'can', 'could', 'would', 'should'
        }
        return common_stopwords

    def extract_from_pdf(self, file_content: bytes) -> Optional[str]:
        """
        Extract text from PDF file

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            text = ""

            for page in pdf_reader.pages:
                text += page.extract_text()

            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting PDF: {str(e)}")
            return None

    def extract_from_docx(self, file_content: bytes) -> Optional[str]:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            docx_file = Document(io.BytesIO(file_content))
            text = ""

            for paragraph in docx_file.paragraphs:
                text += paragraph.text + "\n"

            for table in docx_file.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "

            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            return None

    def clean_text(self, text: str) -> str:
        """
        Clean and preprocess text

        Args:
            text: Raw text from document

        Returns:
            Cleaned text
        """
        # Convert to lowercase
        text = text.lower()

        # Remove special characters but keep spaces
        text = re.sub(r'[^a-z0-9\s\.]', ' ', text)

        # Remove extra whitespace
        text = ' '.join(text.split())

        # Remove numbers-only tokens
        text = re.sub(r'\b\d+\b', '', text)

        # Remove extra spaces again
        text = ' '.join(text.split())

        return text

    def remove_stopwords(self, text: str) -> str:
        """
        Remove common stopwords from text

        Args:
            text: Input text

        Returns:
            Text with stopwords removed
        """
        words = text.split()
        filtered_words = [
            word for word in words if word.lower() not in self.stopwords
        ]
        return ' '.join(filtered_words)

    def extract_and_preprocess(self, file_content: bytes, filename: str) -> Optional[str]:
        """
        Unified method to extract and preprocess document

        Args:
            file_content: File bytes
            filename: Original filename

        Returns:
            Preprocessed text
        """
        if filename.endswith('.pdf'):
            text = self.extract_from_pdf(file_content)
        elif filename.endswith('.docx'):
            text = self.extract_from_docx(file_content)
        else:
            return None

        if text:
            return self.remove_stopwords(text)
        return None
